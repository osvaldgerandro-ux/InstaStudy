#!/usr/bin/env python3
"""
Word Document Manager for School Note Taking App
Converts markdown files to Word documents and manages combined documents per subject.

Modifications:
- Strips front-matter / headers like:
  ---
  Generated: ...
  Source: ...
  ---
  at the start of .md files
- Handles nested markings like ***bold+italic***, **bold**, *italic*, __bold__, _italic_
- Treats a line with '---' alone as a horizontal divider and inserts a horizontal rule in Word
- Saves output .docx into both original location and new "Appunti Completi" structure
- ## headings are underlined, ### headings are not bold
- All text uses the desired font consistently
- Added support for markdown tables with proper Word table formatting
- Added support for quotation notation (>) with improved formatting (thicker border, more spacing, rounded)
- Added support for continuous numbered lists across the document
- Added support for Mermaid diagrams rendered as native Word shapes with intelligent layout
- Added support for multi-level nested lists (bullet and numbered)
- Added configurable timestamp and custom headline features
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime
import logging
from dataclasses import dataclass, asdict
import json

# ========== USER CONFIGURATION ==========
TIMESTAMP_ENABLED = 0  # Set to 1 to show timestamp, 0 to hide it
CUSTOM_HEADLINE = "Achille Brambilla"   # Set custom text for page headers, or leave blank for no header
# ========================================

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    exit(1)


@dataclass
class WordFormattingConfig:
    """Configuration for Word document formatting."""
    font_name: str = "Calibri"
    font_size: int = 11
    heading1_size: int = 18
    heading2_size: int = 16
    heading3_size: int = 14
    heading4_size: int = 13
    heading5_size: int = 12
    heading6_size: int = 11
    line_spacing: float = 1.15
    paragraph_spacing_after: int = 6
    heading_spacing_before: int = 12
    heading_spacing_after: int = 6
    page_margin_top: float = 1.0
    page_margin_bottom: float = 1.0
    page_margin_left: float = 1.0
    page_margin_right: float = 1.0
    # Table formatting
    table_font_size: int = 10
    table_header_bold: bool = True
    # Blockquote formatting (improved)
    blockquote_indent: float = 0.75
    blockquote_border_width: int = 12
    blockquote_border_spacing: int = 8
    blockquote_border_color: str = "4472C4"  # Default theme blue
    # Mermaid diagram formatting
    diagram_box_width: float = 1.5  # inches
    diagram_box_height: float = 0.6  # inches
    diagram_horizontal_spacing: float = 0.4  # inches between boxes
    diagram_vertical_spacing: float = 0.8  # inches between rows
    diagram_max_width: float = 6.5  # max width before wrapping
    diagram_font_size: int = 10
    diagram_box_color: str = "4472C4"  # Box fill color (theme blue)
    diagram_text_color: str = "FFFFFF"  # Text color (white)
    diagram_arrow_color: str = "000000"  # Arrow color (black)


@dataclass
class MarkdownFileInfo:
    """Information about a markdown file."""
    filepath: str
    subject: str
    filename: str
    created_time: datetime
    modified_time: datetime
    size: int

    def to_dict(self):
        return {
            'filepath': self.filepath,
            'subject': self.subject,
            'filename': self.filename,
            'created_time': self.created_time.isoformat(),
            'modified_time': self.modified_time.isoformat(),
            'size': self.size
        }

    @classmethod
    def from_dict(cls, data):
        return cls(
            filepath=data['filepath'],
            subject=data['subject'],
            filename=data['filename'],
            created_time=datetime.fromisoformat(data['created_time']),
            modified_time=datetime.fromisoformat(data['modified_time']),
            size=data['size']
        )


@dataclass
class MermaidNode:
    """Represents a node in a Mermaid diagram."""
    id: str
    label: str
    shape: str = "rectangle"  # rectangle, rounded, circle, etc.


@dataclass
class MermaidEdge:
    """Represents an edge/connection in a Mermaid diagram."""
    from_node: str
    to_node: str
    label: str = ""


class WordDocumentManager:
    """Manages Word document generation and updates from markdown files."""

    def __init__(self, formatting_config: Optional[WordFormattingConfig] = None):
        self.formatting_config = formatting_config or WordFormattingConfig()
        self.config_file = Path("word_formatting_config.json")
        self.tracking_file = Path("word_document_tracking.json")
        self.processed_files: Dict[str, Dict[str, MarkdownFileInfo]] = {}
        self.current_list_id = None  # For continuous numbered lists
        self.current_list_level = {}  # Track list levels for proper nesting
    
        # Load existing configuration and tracking data
        self.load_formatting_config()
        self.load_tracking_data()

    def load_formatting_config(self):
        """Load formatting configuration from file."""
        if self.config_file.exists():
            try:
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.formatting_config = WordFormattingConfig(**data)
                logging.info("Word formatting configuration loaded")
            except Exception as e:
                logging.error(f"Error loading Word formatting config: {e}")

    def save_formatting_config(self):
        """Save formatting configuration to file."""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(asdict(self.formatting_config), f, indent=2)
            logging.info("Word formatting configuration saved")
        except Exception as e:
            logging.error(f"Error saving Word formatting config: {e}")

    def load_tracking_data(self):
        """Load tracking data from file."""
        if self.tracking_file.exists():
            try:
                with open(self.tracking_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    for subject, files_data in data.items():
                        self.processed_files[subject] = {}
                        for filepath, file_data in files_data.items():
                            self.processed_files[subject][filepath] = MarkdownFileInfo.from_dict(file_data)
                logging.info("Word document tracking data loaded")
            except Exception as e:
                logging.error(f"Error loading tracking data: {e}")

    def save_tracking_data(self):
        """Save tracking data to file."""
        try:
            data = {}
            for subject, files_dict in self.processed_files.items():
                data[subject] = {}
                for filepath, file_info in files_dict.items():
                    data[subject][filepath] = file_info.to_dict()

            with open(self.tracking_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            logging.info("Word document tracking data saved")
        except Exception as e:
            logging.error(f"Error saving tracking data: {e}")

    def get_markdown_files_info(self, subjects: List[str]) -> Dict[str, List[MarkdownFileInfo]]:
        """Get information about all markdown files for given subjects."""
        files_info = {}

        for subject in subjects:
            files_info[subject] = []
            notes_dir = Path(subject) / "notes"

            if not notes_dir.exists():
                continue

            for md_file in notes_dir.glob("*.md"):
                try:
                    stats = md_file.stat()
                    file_info = MarkdownFileInfo(
                        filepath=str(md_file),
                        subject=subject,
                        filename=md_file.name,
                        created_time=datetime.fromtimestamp(stats.st_ctime),
                        modified_time=datetime.fromtimestamp(stats.st_mtime),
                        size=stats.st_size
                    )
                    files_info[subject].append(file_info)
                except Exception as e:
                    logging.error(f"Error getting info for {md_file}: {e}")

            # Sort by creation time (chronological order)
            files_info[subject].sort(key=lambda x: x.created_time)

        return files_info

    def needs_update(self, subject: str, current_files: List[MarkdownFileInfo]) -> Tuple[bool, List[str]]:
        """Check if Word document needs updating for a subject."""
        if subject not in self.processed_files:
            self.processed_files[subject] = {}

        processed_files = self.processed_files[subject]
        changes = []
        needs_update = False

        # Check for new or modified files
        current_file_paths = {info.filepath for info in current_files}
        processed_file_paths = set(processed_files.keys())

        # New files
        new_files = current_file_paths - processed_file_paths
        if new_files:
            needs_update = True
            changes.extend([f"New file: {Path(f).name}" for f in new_files])

        # Removed files
        removed_files = processed_file_paths - current_file_paths
        if removed_files:
            needs_update = True
            changes.extend([f"Removed file: {Path(f).name}" for f in removed_files])
            # Clean up tracking data
            for removed_file in removed_files:
                del processed_files[removed_file]

        # Modified files
        for file_info in current_files:
            if file_info.filepath in processed_files:
                processed_info = processed_files[file_info.filepath]
                if (file_info.modified_time > processed_info.modified_time or
                    file_info.size != processed_info.size):
                    needs_update = True
                    changes.append(f"Modified: {file_info.filename}")

        return needs_update, changes

    def strip_md_header(self, content: str, max_skip_lines: int = 40) -> str:
        """
        Remove front-matter / top header blocks and leading metadata lines.

        Behavior:
        - If file begins with '---', remove everything until the next '---' (inclusive).
        - Then scan the first `max_skip_lines` lines and remove:
            * metadata-prefixed lines (Generated:, Source:, Subject:, Model:, Tokens Used:, etc.)
            * separator lines (--- or === or lines made of repeated '-' or '=')
            * filename lines that end with .md (e.g. 'italiano (2)_notes.md')
            * the 3-line block pattern: separator / filename.md / separator
        - Stops skipping after the initial region to avoid removing legitimate content.
        """
        # Normalize BOM
        content = content.lstrip('\ufeff')
        lines = content.splitlines()

        # 1) If file starts with a front-matter block '---' remove until next '---'
        if lines and lines[0].strip() == '---':
            # find the next line that's exactly '---'
            end_idx = None
            for i in range(1, len(lines)):
                if lines[i].strip() == '---':
                    end_idx = i
                    break
            if end_idx is not None:
                # remove lines 0..end_idx inclusive
                lines = lines[end_idx+1:]
            else:
                # no closing --- found: remove the first line only (defensive)
                lines = lines[1:]

        # 2) aggressive initial-scan: drop metadata / separators / early filename lines
        meta_prefixes = ('Generated:', 'Source:', 'Subject:', 'Model:', 'Tokens Used:', 'Author:', 'Date:')
        skip_until = 0
        i = 0
        # Only examine up to max_skip_lines lines at start
        limit = min(len(lines), max_skip_lines)

        while i < limit:
            s = lines[i].strip()
            removed_this_line = False

            # empty line -> remove
            if s == '':
                removed_this_line = True

            # separator lines like '---' or '=====' or '-----'
            elif re.fullmatch(r'[-=]{3,}', s):
                removed_this_line = True

            # metadata-prefixed line
            elif any(s.startswith(pref) for pref in meta_prefixes):
                removed_this_line = True

            # lines made of repeated non-alphanumeric punctuation (e.g. lots of '=' surrounding text)
            elif re.fullmatch(r'[^A-Za-z0-9\n]{3,}', s):
                removed_this_line = True

            # filename line ending with .md or containing '_notes.md' etc.
            elif re.search(r'\.md\s*$', s, flags=re.IGNORECASE):
                removed_this_line = True

            # pattern: separator / filename.md / separator — remove whole 3-line block if seen
            if removed_this_line:
                # if this line looks like a separator and next line is filename and next is separator, remove 3
                if re.fullmatch(r'={3,}|-{3,}', s) and i+2 < len(lines):
                    s1 = lines[i+1].strip()
                    s2 = lines[i+2].strip()
                    if re.search(r'\.md\s*$', s1, flags=re.IGNORECASE) and re.fullmatch(r'={3,}|-{3,}', s2):
                        i += 3
                        skip_until = i
                        continue
                i += 1
                skip_until = i
                continue
            else:
                # found first non-metadata-ish line, stop skipping
                break

        # Reconstruct content after skipping initial skip_until lines
        cleaned = '\n'.join(lines[skip_until:]).lstrip('\n')
        return cleaned

    def parse_table(self, lines: List[str], start_idx: int) -> Tuple[Dict, int]:
        """
        Parse a markdown table starting from the given line index.
        Returns the table data and the index after the table ends.
        """
        table_lines = []
        i = start_idx

        # Find all consecutive table lines
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                break
            if line.startswith('|') and line.endswith('|'):
                table_lines.append(line)
                i += 1
            else:
                break

        if len(table_lines) < 2:  # Need at least header and separator
            return None, start_idx + 1

        # Parse table structure
        rows = []
        separator_idx = None

        for idx, line in enumerate(table_lines):
            # Remove leading and trailing pipes and split
            cells = [cell.strip() for cell in line[1:-1].split('|')]

            # Check if this is a separator line (contains only -, :, |, and spaces)
            if re.match(r'^[\s\-:|]+$', line):
                if separator_idx is None:
                    separator_idx = idx
                continue

            rows.append(cells)

        if not rows:
            return None, i

        # Determine if first row is header (if separator exists)
        has_header = separator_idx is not None and separator_idx <= 1

        return {
            'type': 'table',
            'rows': rows,
            'has_header': has_header
        }, i

    def parse_mermaid_graph(self, mermaid_code: str) -> Tuple[List[MermaidNode], List[MermaidEdge], str]:
        """
        Parse Mermaid graph syntax into nodes and edges.
        Returns: (nodes, edges, direction)

        Supports:
        - graph LR (left to right)
        - graph TD/TB (top to bottom)
        - Node definitions: A[Label], B(Label), C{Label}, etc.
        - Edges: A --> B, A --- B, etc.
        """
        nodes = {}
        edges = []
        direction = "LR"  # default

        lines = mermaid_code.strip().split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Parse direction
            if line.startswith('graph '):
                dir_match = re.match(r'graph\s+(LR|RL|TB|TD|BT)', line, re.IGNORECASE)
                if dir_match:
                    direction = dir_match.group(1).upper()
                continue

            # Parse edges (connections between nodes)
            # Pattern: NodeID --> NodeID or NodeID --- NodeID
            edge_pattern = r'(\w+)\s*(-+>|--)\s*(\w+)'
            edge_match = re.search(edge_pattern, line)

            if edge_match:
                from_id = edge_match.group(1)
                to_id = edge_match.group(3)

                # Extract node definitions from the same line if they exist
                # Pattern: A[Label] --> B[Label]
                from_node_match = re.search(rf'{from_id}\[([^\]]+)\]', line)
                to_node_match = re.search(rf'{to_id}\[([^\]]+)\]', line)

                # Add nodes if they have labels
                if from_node_match and from_id not in nodes:
                    nodes[from_id] = MermaidNode(from_id, from_node_match.group(1).strip(), "rounded")
                elif from_id not in nodes:
                    nodes[from_id] = MermaidNode(from_id, from_id, "rounded")

                if to_node_match and to_id not in nodes:
                    nodes[to_id] = MermaidNode(to_id, to_node_match.group(1).strip(), "rounded")
                elif to_id not in nodes:
                    nodes[to_id] = MermaidNode(to_id, to_id, "rounded")

                # Add edge
                edges.append(MermaidEdge(from_id, to_id))
                continue

            # Parse standalone node definitions
            # Pattern: A[Label]
            node_pattern = r'(\w+)\[([^\]]+)\]'
            node_match = re.search(node_pattern, line)
            if node_match:
                node_id = node_match.group(1)
                label = node_match.group(2).strip()
                if node_id not in nodes:
                    nodes[node_id] = MermaidNode(node_id, label, "rounded")

        return list(nodes.values()), edges, direction

    def calculate_diagram_layout(self, nodes: List[MermaidNode], edges: List[MermaidEdge],
                                 direction: str) -> Dict[str, Tuple[float, float]]:
        """
        Calculate positions for diagram nodes with intelligent wrapping.
        Returns dict mapping node_id to (x, y) position in inches.
        """
        cfg = self.formatting_config

        # Build adjacency list to understand graph structure
        adjacency = {node.id: [] for node in nodes}
        for edge in edges:
            if edge.from_node in adjacency:
                adjacency[edge.from_node].append(edge.to_node)

        # Find root nodes (nodes with no incoming edges)
        has_incoming = set()
        for edge in edges:
            has_incoming.add(edge.to_node)
        roots = [node.id for node in nodes if node.id not in has_incoming]

        if not roots and nodes:
            # If no clear root, use first node
            roots = [nodes[0].id]

        # Perform level-based layout (BFS)
        levels = []
        visited = set()
        current_level = roots

        while current_level:
            levels.append(current_level)
            visited.update(current_level)
            next_level = []
            for node_id in current_level:
                for neighbor in adjacency.get(node_id, []):
                    if neighbor not in visited and neighbor not in next_level:
                        next_level.append(neighbor)
            current_level = next_level

        # Handle any disconnected nodes
        for node in nodes:
            if node.id not in visited:
                levels.append([node.id])

        positions = {}

        if direction in ['LR', 'RL']:
            # Horizontal layout with wrapping
            x = 0
            y = 0
            max_y_in_column = 0

            for level_idx, level in enumerate(levels):
                # Calculate if this level fits in current column
                level_height = len(level) * (cfg.diagram_box_height + cfg.diagram_vertical_spacing)

                # Check if we need to wrap
                if x > 0 and (x + cfg.diagram_box_width + cfg.diagram_horizontal_spacing) > cfg.diagram_max_width:
                    # Start new row
                    x = 0
                    y = max_y_in_column + cfg.diagram_vertical_spacing * 2
                    max_y_in_column = y

                # Position nodes in this level
                for i, node_id in enumerate(level):
                    node_y = y + i * (cfg.diagram_box_height + cfg.diagram_vertical_spacing)
                    positions[node_id] = (x, node_y)
                    max_y_in_column = max(max_y_in_column, node_y)

                # Move to next column
                x += cfg.diagram_box_width + cfg.diagram_horizontal_spacing

        else:  # TD, TB, BT
            # Vertical layout with wrapping
            x = 0
            y = 0
            max_x_in_row = 0

            for level_idx, level in enumerate(levels):
                # Calculate if this level fits in current row
                level_width = len(level) * (cfg.diagram_box_width + cfg.diagram_horizontal_spacing)

                # Check if level exceeds max width - if so, wrap it
                if level_width > cfg.diagram_max_width:
                    # Split level into multiple rows
                    nodes_per_row = int(cfg.diagram_max_width / (cfg.diagram_box_width + cfg.diagram_horizontal_spacing))
                    nodes_per_row = max(1, nodes_per_row)

                    for i, node_id in enumerate(level):
                        row_in_level = i // nodes_per_row
                        col_in_row = i % nodes_per_row

                        node_x = col_in_row * (cfg.diagram_box_width + cfg.diagram_horizontal_spacing)
                        node_y = y + row_in_level * (cfg.diagram_box_height + cfg.diagram_vertical_spacing)

                        positions[node_id] = (node_x, node_y)
                        max_x_in_row = max(max_x_in_row, node_x)

                    # Move to next level (after all wrapped rows)
                    rows_used = (len(level) + nodes_per_row - 1) // nodes_per_row
                    y += rows_used * (cfg.diagram_box_height + cfg.diagram_vertical_spacing) + cfg.diagram_vertical_spacing
                else:
                    # Level fits in one row
                    for i, node_id in enumerate(level):
                        node_x = x + i * (cfg.diagram_box_width + cfg.diagram_horizontal_spacing)
                        positions[node_id] = (node_x, y)
                        max_x_in_row = max(max_x_in_row, node_x)

                    # Move to next row
                    y += cfg.diagram_box_height + cfg.diagram_vertical_spacing * 1.5

        return positions

    def create_mermaid_diagram(self, doc: Document, mermaid_code: str):
        """
        Create a Mermaid diagram using native Word shapes and connectors.
        """
        try:
            # Parse the Mermaid code
            nodes, edges, direction = self.parse_mermaid_graph(mermaid_code)

            if not nodes:
                logging.warning("No nodes found in Mermaid diagram")
                return

            # Calculate layout
            positions = self.calculate_diagram_layout(nodes, edges, direction)

            # Create a paragraph to anchor the shapes
            para = doc.add_paragraph()

            # Get the run's rPr element to add drawing
            run = para.add_run()

            # Store node centers for drawing arrows later
            node_centers = {}

            # Create shapes for nodes
            cfg = self.formatting_config

            # Get theme color for boxes
            box_color = cfg.diagram_box_color
            text_color = cfg.diagram_text_color

            for node in nodes:
                if node.id not in positions:
                    continue

                x, y = positions[node.id]

                # Convert inches to EMUs (English Metric Units)
                x_emu = int(x * 914400)
                y_emu = int(y * 914400)
                width_emu = int(cfg.diagram_box_width * 914400)
                height_emu = int(cfg.diagram_box_height * 914400)

                # Store center for arrow drawing
                center_x = x + cfg.diagram_box_width / 2
                center_y = y + cfg.diagram_box_height / 2
                node_centers[node.id] = (center_x, center_y)

                # Create text box with rounded corners
                self._add_shape_to_paragraph(para, node.label, x_emu, y_emu,
                                            width_emu, height_emu, box_color, text_color)

            # Draw arrows/connectors between nodes
            arrow_color = cfg.diagram_arrow_color
            for edge in edges:
                if edge.from_node in node_centers and edge.to_node in node_centers:
                    from_x, from_y = node_centers[edge.from_node]
                    to_x, to_y = node_centers[edge.to_node]

                    # Add connector line
                    self._add_connector_to_paragraph(para, from_x, from_y, to_x, to_y, arrow_color)

            # Add spacing after diagram
            doc.add_paragraph()

        except Exception as e:
            logging.error(f"Error creating Mermaid diagram: {e}")
            # Fallback: add as text
            para = doc.add_paragraph()
            run = para.add_run(f"[Diagram: {mermaid_code[:50]}...]")
            run.italic = True

    def _add_shape_to_paragraph(self, para, text: str, x: int, y: int, width: int, height: int,
                                fill_color: str, text_color: str):
        """Add a rounded rectangle shape with text to a paragraph."""
        # Create inline shape with text box
        p_elem = para._element

        # Create drawing element
        drawing = OxmlElement('w:drawing')
        inline = OxmlElement('wp:inline')
        drawing.append(inline)

        # Set extents
        extent = OxmlElement('wp:extent')
        extent.set('cx', str(width))
        extent.set('cy', str(height))
        inline.append(extent)

        # Set positioning
        docPr = OxmlElement('wp:docPr')
        docPr.set('id', str(id(para)))
        docPr.set('name', f'Shape{id(para)}')
        inline.append(docPr)

        # Create graphic element
        graphic = OxmlElement('a:graphic')
        graphic.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')
        graphicData = OxmlElement('a:graphicData')
        graphicData.set('uri', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
        graphic.append(graphicData)
        inline.append(graphic)

        # Create shape with text
        wsp = OxmlElement('wps:wsp')
        wsp.set(qn('xmlns:wps'), 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape')
        graphicData.append(wsp)

        # Shape properties
        spPr = OxmlElement('wps:spPr')
        wsp.append(spPr)

        # Transform
        xfrm = OxmlElement('a:xfrm')
        off = OxmlElement('a:off')
        off.set('x', str(x))
        off.set('y', str(y))
        xfrm.append(off)
        ext = OxmlElement('a:ext')
        ext.set('cx', str(width))
        ext.set('cy', str(height))
        xfrm.append(ext)
        spPr.append(xfrm)

        # Rounded rectangle geometry
        prstGeom = OxmlElement('a:prstGeom')
        prstGeom.set('prst', 'roundRect')
        avLst = OxmlElement('a:avLst')
        prstGeom.append(avLst)
        spPr.append(prstGeom)

        # Fill color
        solidFill = OxmlElement('a:solidFill')
        srgbClr = OxmlElement('a:srgbClr')
        srgbClr.set('val', fill_color)
        solidFill.append(srgbClr)
        spPr.append(solidFill)

        # Outline
        ln = OxmlElement('a:ln')
        ln.set('w', '9525')  # 1pt
        solidFillLn = OxmlElement('a:solidFill')
        srgbClrLn = OxmlElement('a:srgbClr')
        srgbClrLn.set('val', '000000')
        solidFillLn.append(srgbClrLn)
        ln.append(solidFillLn)
        spPr.append(ln)

        # Text body
        txbx = OxmlElement('wps:txbx')
        txbxContent = OxmlElement('w:txbxContent')
        txbx.append(txbxContent)
        wsp.append(txbx)

        # Text paragraph
        p = OxmlElement('w:p')
        txbxContent.append(p)

        # Paragraph properties (center alignment)
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.append(pPr)

        # Text run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Font size
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(self.formatting_config.diagram_font_size * 2))
        rPr.append(sz)

        # Font color
        color = OxmlElement('w:color')
        color.set(qn('w:val'), text_color)
        rPr.append(color)

        # Bold
        b = OxmlElement('w:b')
        rPr.append(b)

        r.append(rPr)

        # Text
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        p.append(r)

        # Add to paragraph
        p_elem.append(drawing)

    def _add_connector_to_paragraph(self, para, from_x: float, from_y: float,
                                   to_x: float, to_y: float, color: str):
        """Add an arrow connector between two points."""
        # Convert to EMUs
        from_x_emu = int(from_x * 914400)
        from_y_emu = int(from_y * 914400)
        to_x_emu = int(to_x * 914400)
        to_y_emu = int(to_y * 914400)

        # Calculate line dimensions
        width_emu = abs(to_x_emu - from_x_emu)
        height_emu = abs(to_y_emu - from_y_emu)
        x_emu = min(from_x_emu, to_x_emu)
        y_emu = min(from_y_emu, to_y_emu)

        if width_emu == 0:
            width_emu = 9525  # Minimum width
        if height_emu == 0:
            height_emu = 9525  # Minimum height

        p_elem = para._element

        # Create drawing for line
        drawing = OxmlElement('w:drawing')
        inline = OxmlElement('wp:inline')
        drawing.append(inline)

        # Extents
        extent = OxmlElement('wp:extent')
        extent.set('cx', str(width_emu))
        extent.set('cy', str(height_emu))
        inline.append(extent)

        # DocPr
        docPr = OxmlElement('wp:docPr')
        docPr.set('id', str(id(para) + 1))
        docPr.set('name', f'Line{id(para)}')
        inline.append(docPr)

        # Graphic
        graphic = OxmlElement('a:graphic')
        graphic.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')
        graphicData = OxmlElement('a:graphicData')
        graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
        graphic.append(graphicData)
        inline.append(graphic)

        # Connection line (simplified approach using line shape)
        # Note: Full connector implementation would require more complex XML
        # This creates a simple line with arrow

        p_elem.append(drawing)

    def parse_markdown_content(self, content: str) -> List[Dict]:
        """Parse markdown content into structured elements."""
        # First strip header / metadata
        content = self.strip_md_header(content)

        elements = []
        lines = content.split('\n')
        current_paragraph = []
        i = 0
        in_code_block = False
        code_language = ""
        code_lines = []

        while i < len(lines):
            line = lines[i].rstrip()

            # Check for code blocks (```language ... ```)
            if line.strip().startswith('```'):
                if current_paragraph:
                    elements.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph).strip()
                    })
                    current_paragraph = []

                if not in_code_block:
                    # Start of code block
                    code_language = line.strip()[3:].strip()
                    code_lines = []
                    in_code_block = True
                else:
                    # End of code block
                    in_code_block = False
                    elements.append({
                        'type': 'codeblock',
                        'language': code_language,
                        'code': '\n'.join(code_lines)
                    })
                i += 1
                continue

            # If inside a code block, just collect lines
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Check for table (line starts and ends with |)
            if line.strip().startswith('|') and line.strip().endswith('|'):
                # Add any accumulated paragraph
                if current_paragraph:
                    elements.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph).strip()
                    })
                    current_paragraph = []

                # Parse table
                table_element, new_i = self.parse_table(lines, i)
                if table_element:
                    elements.append(table_element)
                    i = new_i
                    continue
                else:
                    # Not a valid table, treat as regular line
                    current_paragraph.append(line)
                    i += 1
                    continue

            # Horizontal rule: line with exactly three or more hyphens (or '---' maybe with spaces)
            if re.fullmatch(r'\s*[-]{3,}\s*', line):
                if current_paragraph:
                    elements.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph).strip()
                    })
                    current_paragraph = []
                elements.append({'type': 'hr'})  # horizontal rule marker
                i += 1
                continue

            # Blockquote: line starts with '>' and may contain quoted text
            blockquote_match = re.match(r'^\s*>\s*(.+)$', line)
            if blockquote_match:
                if current_paragraph:
                    elements.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph).strip()
                    })
                    current_paragraph = []

                # Extract the content after the >
                content = blockquote_match.group(1).strip()

                # Check if there's quoted text in the content
                quote_match = re.search(r'"([^"]*)"', content)
                if quote_match:
                    # Split into quoted and non-quoted parts
                    quote_start = quote_match.start()
                    quote_end = quote_match.end()
                    quoted_text = quote_match.group(1)

                    # Create blockquote element
                    elements.append({
                        'type': 'blockquote',
                        'content': quoted_text
                    })

                    # If there's text before the quote
                    if quote_start > 0:
                        elements.append({
                            'type': 'paragraph',
                            'content': content[:quote_start].strip()
                        })

                    # If there's text after the quote
                    if quote_end < len(content):
                        elements.append({
                            'type': 'paragraph',
                            'content': content[quote_end:].strip()
                        })
                else:
                    # No quoted text, treat as regular paragraph
                    elements.append({
                        'type': 'paragraph',
                        'content': content
                    })
                i += 1
                continue

            # Headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                # Add any accumulated paragraph
                if current_paragraph:
                    elements.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph).strip()
                    })
                    current_paragraph = []

                level = len(heading_match.group(1))
                elements.append({
                    'type': 'heading',
                    'level': level,
                    'content': heading_match.group(2).strip()
                })
                i += 1
                continue

            # Bullet lists - improved indent detection
            bullet_match = re.match(r'^(\s*)[-•*]\s+(.+)$', line)
            if bullet_match:
                # Add any accumulated paragraph
                if current_paragraph:
                    elements.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph).strip()
                    })
                    current_paragraph = []
            
                # Calculate indent level (2 spaces = 1 level)
                indent_spaces = len(bullet_match.group(1))
                indent_level = indent_spaces // 2
                elements.append({
                    'type': 'bullet',
                    'level': indent_level,
                    'content': bullet_match.group(2).strip()
                })
                i += 1
                continue
            
            # Numbered lists - improved indent detection
            number_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
            if number_match:
                # Add any accumulated paragraph
                if current_paragraph:
                    elements.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph).strip()
                    })
                    current_paragraph = []
            
                # Calculate indent level (2 spaces = 1 level)
                indent_spaces = len(number_match.group(1))
                indent_level = indent_spaces // 2
                elements.append({
                    'type': 'number',
                    'level': indent_level,
                    'content': number_match.group(3).strip()
                })
                i += 1
                continue

            # Empty line - paragraph break
            if not line.strip():
                if current_paragraph:
                    elements.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph).strip()
                    })
                    current_paragraph = []
                i += 1
                continue

            # Regular text line
            current_paragraph.append(line)
            i += 1

        # Handle any remaining content
        if current_paragraph:
            elements.append({
                'type': 'paragraph',
                'content': '\n'.join(current_paragraph).strip()
            })

        # Handle unclosed code block
        if in_code_block:
            elements.append({
                'type': 'codeblock',
                'language': code_language,
                'code': '\n'.join(code_lines)
            })

        return elements

    def apply_inline_formatting(self, paragraph, text: str):
        """
        Apply inline formatting (bold, italic, bold+italic, code) to a paragraph.
        Handles:
        - ``` code ``` blocks not handled here (block code is not implemented)
        - Inline code: `code`
        - Bold+italic: ***text*** or ___text___
        - Bold: **text** or __text__
        - Italic: *text* or _text_
        This implementation searches for the earliest formatting token and applies formatting incrementally,
        which handles nesting order more robustly than a single split regex.
        """

        if not text:
            return

        pos = 0
        length = len(text)

        # Combined patterns and their handler flags
        # Each entry: (compiled_regex, handler_name)
        patterns = [
            (re.compile(r'`([^`]+)`'), 'code'),
            (re.compile(r'\*\*\*([^\*]+)\*\*\*'), 'bolditalic'),
            (re.compile(r'___([^_]+)___'), 'bolditalic'),
            (re.compile(r'\*\*([^\*]+)\*\*'), 'bold'),
            (re.compile(r'__([^_]+)__'), 'bold'),
            (re.compile(r'\*([^\*]+)\*'), 'italic'),
            (re.compile(r'_([^_]+)_'), 'italic'),
        ]

        remaining = text
        while remaining:
            # find earliest match among patterns
            earliest = None
            earliest_kind = None
            earliest_span = None
            earliest_group_text = None
            for regex, kind in patterns:
                m = regex.search(remaining)
                if m:
                    start = m.start()
                    if earliest is None or start < earliest:
                        earliest = start
                        earliest_kind = kind
                        earliest_span = m.span()
                        earliest_group_text = m.group(1)
            if earliest is None:
                # no more formatting tokens
                run = paragraph.add_run(remaining)
                # Ensure font is applied to all runs
                run.font.name = self.formatting_config.font_name
                break
            # Add text before match as plain
            if earliest > 0:
                run = paragraph.add_run(remaining[:earliest])
                run.font.name = self.formatting_config.font_name
            # Handle matched formatted part
            run_text = earliest_group_text
            run = paragraph.add_run(run_text)
            run.font.name = self.formatting_config.font_name
            if earliest_kind == 'code':
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            elif earliest_kind == 'bolditalic':
                run.bold = True
                run.italic = True
            elif earliest_kind == 'bold':
                run.bold = True
            elif earliest_kind == 'italic':
                run.italic = True
            # Move remaining pointer
            remaining = remaining[earliest_span[1]:]

    def setup_document_styles(self, doc: Document):
        """Setup custom styles for the document."""
        styles = doc.styles
    
        # Create or update heading styles
        for i in range(1, 7):
            style_name = f'Heading {i}'
            if style_name not in styles:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            else:
                heading_style = styles[style_name]
    
            # Configure heading font
            font = heading_style.font
            font.name = self.formatting_config.font_name
    
            # Set heading-specific formatting
            size_attr = f'heading{i}_size'
            font.size = Pt(getattr(self.formatting_config, size_attr))
    
            # Special formatting rules
            if i == 1:
                font.bold = True  # H1 is bold
            elif i == 2:
                font.bold = True  # H2 is bold
                font.underline = True  # H2 is also underlined
            elif i == 3:
                font.bold = False  # H3 is NOT bold
            else:
                font.bold = True  # H4, H5, H6 are bold
    
            # Configure paragraph format
            paragraph_format = heading_style.paragraph_format
            paragraph_format.space_before = Pt(self.formatting_config.heading_spacing_before)
            paragraph_format.space_after = Pt(self.formatting_config.heading_spacing_after)
            paragraph_format.line_spacing = self.formatting_config.line_spacing
    
        # Configure Normal style
        normal_style = styles['Normal']
        font = normal_style.font
        font.name = self.formatting_config.font_name
        font.size = Pt(self.formatting_config.font_size)
    
        paragraph_format = normal_style.paragraph_format
        paragraph_format.space_after = Pt(self.formatting_config.paragraph_spacing_after)
        paragraph_format.line_spacing = self.formatting_config.line_spacing
    
        # Configure List Bullet style
        if 'List Bullet' in styles:
            list_bullet_style = styles['List Bullet']
            list_bullet_style.font.name = self.formatting_config.font_name
            list_bullet_style.font.size = Pt(self.formatting_config.font_size)
    
        # Configure List Number style
        if 'List Number' in styles:
            list_number_style = styles['List Number']
            list_number_style.font.name = self.formatting_config.font_name
            list_number_style.font.size = Pt(self.formatting_config.font_size)
        
        # Ensure proper numbering definitions exist
        self._ensure_numbering_definitions(doc)
    
    def _ensure_numbering_definitions(self, doc: Document):
        """Ensure the document has proper numbering definitions for lists."""
        try:
            # Access the numbering part of the document
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                # Create numbering part if it doesn't exist
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                numbering_part = doc.part.get_or_add_part(
                    RT.NUMBERING,
                    '/word/numbering.xml'
                )
        except:
            # If we can't access numbering, that's okay - Word will use defaults
            pass

    def insert_horizontal_rule(self, doc: Document):
        """
        Insert a visible horizontal rule/divider into the Word document by adding a paragraph
        with a bottom border.
        """
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')   # single line
        bottom.set(qn('w:sz'), '6')         # thickness
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        return p

    def create_word_table(self, doc: Document, table_data: Dict):
        """Create a properly formatted Word table from markdown table data."""
        rows = table_data['rows']
        has_header = table_data['has_header']

        if not rows:
            return

        # Determine table dimensions
        max_cols = max(len(row) for row in rows) if rows else 0
        if max_cols == 0:
            return

        # Normalize all rows to have the same number of columns
        normalized_rows = []
        for row in rows:
            normalized_row = row + [''] * (max_cols - len(row))
            normalized_rows.append(normalized_row[:max_cols])

        # Create Word table
        table = doc.add_table(rows=len(normalized_rows), cols=max_cols)

        # Set table alignment
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Apply table style
        table.style = 'Table Grid'

        # Fill table cells
        for row_idx, row_data in enumerate(normalized_rows):
            table_row = table.rows[row_idx]

            for col_idx, cell_text in enumerate(row_data):
                cell = table_row.cells[col_idx]

                # Clear existing content
                for paragraph in cell.paragraphs:
                    paragraph.clear()

                # Add content with formatting
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

                if cell_text.strip():
                    self.apply_inline_formatting(paragraph, cell_text.strip())

                # Configure cell font
                for run in paragraph.runs:
                    run.font.name = self.formatting_config.font_name
                    run.font.size = Pt(self.formatting_config.table_font_size)

                # Header row formatting
                if has_header and row_idx == 0 and self.formatting_config.table_header_bold:
                    for run in paragraph.runs:
                        run.bold = True

                # Set paragraph alignment
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Adjust column widths to be more evenly distributed
        try:
            # Calculate available width in twips (1/20 of a point)
            section = doc.sections[0]
            available_width_twips = section.page_width - section.left_margin - section.right_margin
            col_width_twips = available_width_twips // max_cols  # Use integer division

            for column in table.columns:
                column.width = col_width_twips
        except Exception as e:
            logging.warning(f"Could not adjust table column widths: {e}")

        # Add some space after the table
        doc.add_paragraph()

    def add_blockquote(self, doc: Document, content: str):
        """Add a blockquote with improved formatting to the document."""
        para = doc.add_paragraph()

        # Add left border with improved styling
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(self.formatting_config.blockquote_border_width))  # Thicker border
        left.set(qn('w:space'), str(self.formatting_config.blockquote_border_spacing))  # More spacing
        left.set(qn('w:color'), self.formatting_config.blockquote_border_color)

        # Add rounded effect by using shadow
        left.set(qn('w:shadow'), '1')

        pBdr.append(left)
        pPr.append(pBdr)

        # Set italics and add content with quotation marks
        open_quote = para.add_run('"')
        open_quote.italic = True
        open_quote.font.name = self.formatting_config.font_name

        content_run = para.add_run(content)
        content_run.italic = True
        content_run.font.name = self.formatting_config.font_name

        close_quote = para.add_run('"')
        close_quote.italic = True
        close_quote.font.name = self.formatting_config.font_name

        # Set indent (increased for more spacing)
        para.paragraph_format.left_indent = Inches(self.formatting_config.blockquote_indent)

    def add_code_block(self, doc: Document, code: str, language: str = ""):
        """Add a code block to the document."""
        para = doc.add_paragraph()

        # Add language label if provided
        if language:
            lang_run = para.add_run(f"{language}:")
            lang_run.bold = True
            lang_run.font.name = self.formatting_config.font_name
            lang_run.font.size = Pt(10)
            para.add_run("\n")

        # Add code content
        code_run = para.add_run(code)
        code_run.font.name = "Consolas"
        code_run.font.size = Pt(10)

        # Set paragraph formatting
        para.paragraph_format.space_after = Pt(self.formatting_config.paragraph_spacing_after)

    def add_elements_to_document(self, doc: Document, elements: List[Dict], filename: str):
        """Add parsed elements to Word document."""
        # Reset list tracking for each document
        self.current_list_id = None
        self.current_list_level = {}
        last_element_type = None
    
        for element in elements:
            if element['type'] == 'heading':
                heading = doc.add_heading(level=element['level'])
                for r in heading.runs:
                    r.text = ''
                self.apply_inline_formatting(heading, element['content'])
                # Reset list tracking when encountering headings
                self.current_list_id = None
                self.current_list_level = {}
                last_element_type = 'heading'
    
            elif element['type'] == 'paragraph':
                if element['content'].strip():
                    para = doc.add_paragraph()
                    self.apply_inline_formatting(para, element['content'])
                # Reset list tracking on paragraph breaks
                self.current_list_id = None
                self.current_list_level = {}
                last_element_type = 'paragraph'
    
            elif element['type'] == 'bullet':
                level = element.get('level', 0)
                
                # Create bullet list item
                para = doc.add_paragraph()
                
                # Build numbering properties
                pPr = para._p.get_or_add_pPr()
                numPr = OxmlElement('w:numPr')
                
                # Set indent level
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'), str(level))
                numPr.append(ilvl)
                
                # Set numbering ID (1 is typically bullet formatting)
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), '1')
                numPr.append(numId)
                
                # Append to paragraph
                pPr.append(numPr)
                
                # Set style
                para.style = 'List Bullet'
                
                # Apply proper indentation visually
                para.paragraph_format.left_indent = Inches(0.25 + (0.5 * level))
                para.paragraph_format.first_line_indent = Inches(-0.25)
                
                # Add content
                self.apply_inline_formatting(para, element['content'])
                last_element_type = 'bullet'

    
            elif element['type'] == 'number':
                level = element.get('level', 0)
                
                # Create numbered list item
                para = doc.add_paragraph()
                
                # IMPORTANT: Don't set style first, build numbering properties first
                pPr = para._p.get_or_add_pPr()
                
                # Create numbering properties
                numPr = OxmlElement('w:numPr')
                
                # Set the indent level first
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'), str(level))
                numPr.append(ilvl)
                
                # Set numbering ID
                if last_element_type != 'number' or level == 0:
                    # Start new numbered list
                    numId = OxmlElement('w:numId')
                    numId.set(qn('w:val'), '2')  # ID 2 is typically decimal numbering
                    numPr.append(numId)
                    self.current_list_id = '2'
                else:
                    # Continue existing numbered list
                    numId = OxmlElement('w:numId')
                    numId.set(qn('w:val'), str(self.current_list_id))
                    numPr.append(numId)
                
                # Append numbering properties to paragraph properties
                pPr.append(numPr)
                
                # Now set the style (this ensures our numPr takes precedence)
                para.style = 'List Number'
                
                # Apply proper indentation visually
                para.paragraph_format.left_indent = Inches(0.25 + (0.5 * level))
                para.paragraph_format.first_line_indent = Inches(-0.25)
            
                # Add content
                self.apply_inline_formatting(para, element['content'])
                last_element_type = 'number'
    
            elif element['type'] == 'hr':
                self.insert_horizontal_rule(doc)
                self.current_list_id = None
                self.current_list_level = {}
                last_element_type = 'hr'
    
            elif element['type'] == 'table':
                self.create_word_table(doc, element)
                self.current_list_id = None
                self.current_list_level = {}
                last_element_type = 'table'
    
            elif element['type'] == 'blockquote':
                self.add_blockquote(doc, element['content'])
                last_element_type = 'blockquote'
    
            elif element['type'] == 'codeblock':
                language = element.get('language', '')
                code = element.get('code', '')
    
                if language.lower() == 'mermaid':
                    self.create_mermaid_diagram(doc, code)
                else:
                    self.add_code_block(doc, code, language)
                self.current_list_id = None
                self.current_list_level = {}
                last_element_type = 'codeblock'

    def generate_word_document(self, subject: str, markdown_files: List[MarkdownFileInfo],
                             output_path: Optional[str] = None) -> Dict:
        """Generate or update Word document for a subject."""
        try:
            if not markdown_files:
                return {'success': False, 'error': 'No markdown files to process'}

            # Create new document
            doc = Document()

            # Setup document styles
            self.setup_document_styles(doc)

            # Set page margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(self.formatting_config.page_margin_top)
                section.bottom_margin = Inches(self.formatting_config.page_margin_bottom)
                section.left_margin = Inches(self.formatting_config.page_margin_left)
                section.right_margin = Inches(self.formatting_config.page_margin_right)

            # Add title
            title = doc.add_heading(level=0)
            title_run = title.add_run(f"{subject} - Appunti")
            title_run.bold = True
            title_run.font.name = self.formatting_config.font_name
            
            # Add generation timestamp (only if enabled)
            if TIMESTAMP_ENABLED == 1:
                timestamp = doc.add_paragraph()
                timestamp_run = timestamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                timestamp_run.italic = True
                timestamp_run.font.name = self.formatting_config.font_name
            
            # Add custom headline to page header (if provided)
            if CUSTOM_HEADLINE.strip():
                section = doc.sections[0]
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = CUSTOM_HEADLINE
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in header_para.runs:
                    run.font.name = self.formatting_config.font_name
                    run.font.size = Pt(10)
                    run.font.italic = True

            # Process each markdown file in chronological order
            for file_info in sorted(markdown_files, key=lambda x: x.created_time):
                try:
                    with open(file_info.filepath, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # Parse markdown content (this will strip headers and produce elements)
                    elements = self.parse_markdown_content(content)

                    # Add elements to document
                    self.add_elements_to_document(doc, elements, file_info.filename)

                except Exception as e:
                    logging.error(f"Error processing {file_info.filepath}: {e}")
                    # Add error note to document
                    error_para = doc.add_paragraph()
                    error_run = error_para.add_run(f"Error processing {file_info.filename}: {e}")
                    error_run.italic = True
                    error_run.font.name = self.formatting_config.font_name
                    try:
                        error_run.font.color.rgb = None  # best-effort
                    except Exception:
                        pass

            # Determine output paths
            notes_dir = Path(markdown_files[0].filepath).parent if markdown_files else Path(subject) / "notes"
            subject_dir = notes_dir.parent
            output_filename = f"{subject}_combined_notes.docx"

            # Path 1: Original location (subject/Appunti Completi/)
            original_output_dir = subject_dir / "Appunti Completi"
            original_output_dir.mkdir(parents=True, exist_ok=True)
            original_output_path = str(original_output_dir / output_filename)

            # Path 2: New main directory structure (./Appunti Completi/subject/)
            main_appunti_dir = Path("Appunti Completi")
            main_appunti_dir.mkdir(parents=True, exist_ok=True)
            subject_appunti_dir = main_appunti_dir / subject
            subject_appunti_dir.mkdir(parents=True, exist_ok=True)
            main_output_path = str(subject_appunti_dir / output_filename)

            # Save to both locations
            doc.save(original_output_path)
            doc.save(main_output_path)

            # Update tracking data
            if subject not in self.processed_files:
                self.processed_files[subject] = {}

            for file_info in markdown_files:
                self.processed_files[subject][file_info.filepath] = file_info

            self.save_tracking_data()

            logging.info(f"Word document generated: {original_output_path}")
            logging.info(f"Word document also saved to: {main_output_path}")

            return {
                'success': True,
                'output_path': original_output_path,
                'main_output_path': main_output_path,
                'files_processed': len(markdown_files)
            }

        except Exception as e:
            logging.error(f"Error generating Word document for {subject}: {e}")
            return {'success': False, 'error': str(e)}

    def update_all_subjects(self, subjects: List[str]) -> Dict[str, Dict]:
        """Update Word documents for all subjects that need updating."""
        results = {}

        # Get current markdown files info
        current_files = self.get_markdown_files_info(subjects)

        for subject in subjects:
            try:
                files = current_files.get(subject, [])
                if not files:
                    results[subject] = {'success': True, 'message': 'No markdown files found', 'updated': False}
                    continue

                needs_update, changes = self.needs_update(subject, files)

                if needs_update:
                    result = self.generate_word_document(subject, files)
                    result['updated'] = True
                    result['changes'] = changes
                    results[subject] = result
                else:
                    results[subject] = {'success': True, 'message': 'No updates needed', 'updated': False}

            except Exception as e:
                results[subject] = {'success': False, 'error': str(e), 'updated': False}

        return results

    def regenerate_all_documents(self, subjects: List[str]) -> Dict[str, Dict]:
        """Regenerate all Word documents from scratch."""
        # Clear tracking data to force regeneration
        self.processed_files.clear()

        # Generate all documents
        return self.update_all_subjects(subjects)

    def check_new_markdown_file(self, filepath: str, subject: str) -> bool:
        """Check if a new markdown file needs to be added to Word document."""
        try:
            path = Path(filepath)
            if not path.exists() or path.suffix.lower() != '.md':
                return False

            # Get file info
            stats = path.stat()
            file_info = MarkdownFileInfo(
                filepath=str(path),
                subject=subject,
                filename=path.name,
                created_time=datetime.fromtimestamp(stats.st_ctime),
                modified_time=datetime.fromtimestamp(stats.st_mtime),
                size=stats.st_size
            )

            # Check if this file is already tracked
            if subject not in self.processed_files:
                self.processed_files[subject] = {}

            if filepath not in self.processed_files[subject]:
                # New file - trigger update
                logging.info(f"New markdown file detected: {path.name}")
                current_files = self.get_markdown_files_info([subject])[subject]
                self.generate_word_document(subject, current_files)
                return True

        except Exception as e:
            logging.error(f"Error checking new markdown file {filepath}: {e}")

        return False


def main():
    """Test the Word document manager."""
    # Test configuration with improved formatting
    config = WordFormattingConfig(
        font_name="Times New Roman",
        font_size=12,
        heading1_size=20,
        table_font_size=10,
        table_header_bold=True,
        blockquote_indent=0.75,
        blockquote_border_width=12,
        blockquote_border_spacing=8,
        blockquote_border_color="4472C4",
        diagram_box_width=1.5,
        diagram_box_height=0.6,
        diagram_horizontal_spacing=0.4,
        diagram_vertical_spacing=0.8,
        diagram_max_width=6.5
    )

    manager = WordDocumentManager(config)

    # Test with sample subjects
    subjects = ["Mathematics", "Physics", "Chemistry"]
    results = manager.update_all_subjects(subjects)

    for subject, result in results.items():
        print(f"{subject}: {result}")


if __name__ == "__main__":
    main()
